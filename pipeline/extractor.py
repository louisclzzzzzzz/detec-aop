"""
Extraction de texte depuis les documents (PDF, DOCX, DOC, XLSX, XLS).

Stratégie par type :
  PDF natif  → PyMuPDF (extraction directe)
  PDF scanné → Tesseract OCR
  DOCX       → python-docx (paragraphes + tableaux)
  DOC        → docx2txt (format binaire .doc)
  XLSX       → openpyxl (valeurs des cellules)
  XLS        → xlrd (format binaire .xls)
"""

from pathlib import Path

import fitz  # PyMuPDF
import pytesseract
from PIL import Image
from rich.console import Console

from concurrent.futures import ThreadPoolExecutor

from config import TESSERACT_LANG, TESSERACT_DPI, OCR_MAX_PAGES, MAX_WORKERS
from models.schemas import ExtractedDocument, ExtractionMethod
from utils.pdf_utils import is_scanned_pdf

console = Console()


def _ocr_pdf_tesseract(pdf_path: str) -> tuple[str, int]:
    doc = fitz.open(pdf_path)
    pages_text = []

    for page_num, page in enumerate(doc):
        if page_num >= OCR_MAX_PAGES:
            break
        mat = fitz.Matrix(TESSERACT_DPI / 72, TESSERACT_DPI / 72)
        pix = page.get_pixmap(matrix=mat)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        text = pytesseract.image_to_string(img, lang=TESSERACT_LANG)
        pages_text.append(text)

    processed_pages = min(len(doc), OCR_MAX_PAGES)
    doc.close()
    return "\n".join(pages_text), processed_pages


def _extract_docx(path: Path) -> tuple[str, int]:
    """Extrait le texte d'un fichier .docx via python-docx."""
    import docx
    doc = docx.Document(str(path))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts), len(doc.paragraphs)


def _extract_doc(path: Path) -> tuple[str, int]:
    """
    Extrait le texte d'un fichier .doc.
    Essaie d'abord docx2txt (pour les .doc qui sont en réalité des .docx renommés),
    puis tombe sur une extraction OLE pour les vrais binaires Word 97-2003.
    """
    import re

    # Tentative docx2txt (format zip déguisé en .doc)
    try:
        import docx2txt
        text = docx2txt.process(str(path)) or ""
        lines = [l for l in text.splitlines() if l.strip()]
        if lines:
            return "\n".join(lines), len(lines)
    except Exception:
        pass

    # Fallback OLE (vrai binaire Word 97-2003)
    import olefile
    ole = olefile.OleFileIO(str(path))
    raw = b""
    for stream in ("WordDocument", "1Table", "0Table"):
        if ole.exists(stream):
            raw += ole.openstream(stream).read()
    ole.close()

    text = raw.decode("latin-1", errors="ignore")
    # Extraire les séquences de caractères imprimables (latin + ASCII)
    chunks = re.findall(r'[\x20-\x7E\x80-\xFF]{4,}', text)
    # Filtrer les chunks qui ne sont que des caractères non-texte répétés
    readable = [c for c in chunks if sum(1 for ch in c if ch.isalpha()) > len(c) * 0.3]
    result = " ".join(readable)
    lines = [l for l in result.splitlines() if l.strip()]
    return result, len(lines)


def _extract_xlsx(path: Path) -> tuple[str, int]:
    """Extrait le texte d'un fichier .xlsx via openpyxl."""
    import openpyxl
    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    parts = []
    total_rows = 0
    for sheet in wb.worksheets:
        parts.append(f"[Feuille: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None and str(c).strip()]
            if cells:
                parts.append(" | ".join(cells))
                total_rows += 1
    wb.close()
    return "\n".join(parts), total_rows


def _extract_xls(path: Path) -> tuple[str, int]:
    """Extrait le texte d'un fichier .xls binaire via xlrd."""
    import xlrd
    wb = xlrd.open_workbook(str(path))
    parts = []
    total_rows = 0
    for sheet in wb.sheets():
        parts.append(f"[Feuille: {sheet.name}]")
        for row_idx in range(sheet.nrows):
            cells = [str(sheet.cell_value(row_idx, c)) for c in range(sheet.ncols)
                     if str(sheet.cell_value(row_idx, c)).strip()]
            if cells:
                parts.append(" | ".join(cells))
                total_rows += 1
    return "\n".join(parts), total_rows


def extract_document(doc_path: Path, subdirectory: str = "") -> ExtractedDocument:
    """
    Extrait le texte d'un document en choisissant la stratégie selon l'extension.

    Args:
        doc_path: Chemin vers le fichier.
        subdirectory: Chemin relatif du sous-dossier.

    Returns:
        ExtractedDocument avec le texte extrait et les métadonnées.
    """
    suffix = doc_path.suffix.lower()

    try:
        if suffix == ".pdf":
            is_scanned, native_text, num_pages = is_scanned_pdf(str(doc_path))
            if not is_scanned:
                console.print(f"    📄 [green]PDF natif[/green] ({num_pages} pages)")
                return ExtractedDocument(
                    path=doc_path,
                    filename=doc_path.name,
                    text=native_text,
                    num_pages=num_pages,
                    extraction_method=ExtractionMethod.NATIVE,
                    subdirectory=subdirectory,
                )
            console.print(f"    🔍 [yellow]PDF scanné[/yellow] → OCR ({num_pages} pages)")
            ocr_text, num_pages = _ocr_pdf_tesseract(str(doc_path))
            return ExtractedDocument(
                path=doc_path,
                filename=doc_path.name,
                text=ocr_text,
                num_pages=num_pages,
                extraction_method=ExtractionMethod.OCR_MISTRAL,
                subdirectory=subdirectory,
            )

        elif suffix == ".docx":
            text, num_pages = _extract_docx(doc_path)
            console.print(f"    📝 [cyan]DOCX[/cyan] ({num_pages} paragraphes)")
            return ExtractedDocument(
                path=doc_path,
                filename=doc_path.name,
                text=text,
                num_pages=num_pages,
                extraction_method=ExtractionMethod.DOCX,
                subdirectory=subdirectory,
            )

        elif suffix == ".doc":
            text, num_pages = _extract_doc(doc_path)
            console.print(f"    📝 [cyan]DOC[/cyan] ({num_pages} lignes)")
            return ExtractedDocument(
                path=doc_path,
                filename=doc_path.name,
                text=text,
                num_pages=num_pages,
                extraction_method=ExtractionMethod.DOC,
                subdirectory=subdirectory,
            )

        elif suffix == ".xlsx":
            text, num_pages = _extract_xlsx(doc_path)
            console.print(f"    📊 [magenta]XLSX[/magenta] ({num_pages} lignes)")
            return ExtractedDocument(
                path=doc_path,
                filename=doc_path.name,
                text=text,
                num_pages=num_pages,
                extraction_method=ExtractionMethod.XLSX,
                subdirectory=subdirectory,
            )

        elif suffix == ".xls":
            text, num_pages = _extract_xls(doc_path)
            console.print(f"    📊 [magenta]XLS[/magenta] ({num_pages} lignes)")
            return ExtractedDocument(
                path=doc_path,
                filename=doc_path.name,
                text=text,
                num_pages=num_pages,
                extraction_method=ExtractionMethod.XLS,
                subdirectory=subdirectory,
            )

        else:
            raise ValueError(f"Extension non supportée : {suffix}")

    except Exception as e:
        error_msg = f"{type(e).__name__}: {e}"
        console.print(f"    💥 [red]Erreur[/red] : {error_msg}")
        return ExtractedDocument(
            path=doc_path,
            filename=doc_path.name,
            text="",
            num_pages=0,
            extraction_method=ExtractionMethod.FAILED,
            subdirectory=subdirectory,
            extraction_error=error_msg,
        )


def extract_all(doc_paths: list[Path], root_dir: Path | None = None) -> list[ExtractedDocument]:
    """
    Extrait le texte de tous les documents fournis en parallèle.

    Args:
        doc_paths: Liste de chemins vers les fichiers à traiter.
        root_dir: Dossier racine pour calculer les sous-dossiers relatifs.

    Returns:
        Liste de ExtractedDocument dans le même ordre que doc_paths.
    """
    from utils.file_utils import get_subdirectory

    total = len(doc_paths)

    def _extract_one(item: tuple[int, Path]) -> ExtractedDocument:
        i, doc_path = item
        subdir = get_subdirectory(doc_path, root_dir) if root_dir else ""
        display_path = f"{subdir}/{doc_path.name}" if subdir else doc_path.name
        console.print(f"  [{i + 1}/{total}] [bold]{display_path}[/bold]")
        return extract_document(doc_path, subdirectory=subdir)

    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        return list(executor.map(_extract_one, enumerate(doc_paths)))
